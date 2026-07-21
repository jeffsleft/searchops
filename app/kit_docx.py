"""
W2 — Application Kit brief (.docx).

Renders the kit's *review* artifact: recommended angle, Layer 2 evidence, mismatches to
address, tailored resume bullets, cover-letter hooks, and the full cover-letter text.

This is deliberately NOT a merged resume+cover-letter file. Those are two documents an
employer receives separately; fusing them into one .docx produces something that has to be
taken apart before it can be sent. The send artifacts keep their own routes
(`build_resume_docx` / `build_cover_letter_docx` in app/resume_docx.py, rendered from the
locked docxtpl masters). This brief is what Jeff reads, annotates, and works from.

Built with python-docx directly rather than a locked docxtpl master: the masters exist so
candidate-facing documents can be restyled in Word without touching code
(docs/customize-doc-templates.md). The brief is internal — it never leaves the building —
so it doesn't earn a locked template.
"""
from __future__ import annotations

import io

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

_INK = RGBColor(0x1A, 0x1A, 0x1A)
_MUTED = RGBColor(0x66, 0x66, 0x66)


def _heading(doc, text: str, size: int = 13) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    run.font.color.rgb = _INK


def _muted(doc, text: str, size: int = 8.5) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = _MUTED


def _table(doc, headers: list[str], rows: list[list[str]], widths: list[float] | None = None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for cell, label in zip(table.rows[0].cells, headers):
        cell.text = ""
        run = cell.paragraphs[0].add_run(label)
        run.bold = True
        run.font.size = Pt(8.5)
    for row in rows:
        cells = table.add_row().cells
        for cell, val in zip(cells, row):
            cell.text = ""
            run = cell.paragraphs[0].add_run(val or "—")
            run.font.size = Pt(9)
    return table


def build_kit_docx(kit: dict) -> bytes:
    """
    Return .docx bytes for the Application Kit brief.

    `kit` is the dict from `app.services.kit_service.build_kit_data`. Tolerates missing
    sections — a kit with no research or no letter still produces a usable brief.
    """
    job = kit.get("job") or {}
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    title = doc.add_paragraph()
    title_run = title.add_run(job.get("job_title") or "Application Kit")
    title_run.bold = True
    title_run.font.size = Pt(18)

    subtitle_bits = [b for b in (job.get("company"), job.get("location")) if b]
    if subtitle_bits:
        _muted(doc, "  ·  ".join(subtitle_bits), size=10)

    score = job.get("final_score")
    meta = []
    if score is not None:
        meta.append(f"Score {float(score):.1f} — {job.get('tier', '')}".strip(" —"))
    if job.get("url"):
        meta.append(job["url"])
    if meta:
        _muted(doc, "  |  ".join(meta))

    corpus = kit.get("corpus") or {}
    if corpus.get("is_example"):
        _muted(
            doc,
            "EXAMPLE DATA — built from the fictional demo corpus "
            "(data/Accomplishments_Inventory.example.docx), not a real inventory.",
            size=9,
        )
    elif not corpus.get("available"):
        _muted(
            doc,
            "WARNING: no accomplishments corpus loaded — Layer 2 contributed 0.0 and the "
            "evidence, bullets, and hooks below are empty as a result.",
            size=9,
        )

    if kit.get("angle"):
        _heading(doc, "Recommended angle")
        p = doc.add_paragraph()
        run = p.add_run(f"“{kit['angle']}”")
        run.italic = True
        run.font.size = Pt(10.5)

    evidence = kit.get("evidence") or []
    if evidence:
        _heading(doc, f"Evidence ({len(evidence)})")
        _table(
            doc,
            ["JD requirement", "Matched accomplishment", "Strength"],
            [
                [
                    str(e.get("jd_requirement") or ""),
                    str(e.get("matched_accomplishment") or ""),
                    str(e.get("strength") or ""),
                ]
                for e in evidence
                if isinstance(e, dict)
            ],
        )

    mismatches = kit.get("mismatches") or []
    if mismatches:
        _heading(doc, f"Address proactively ({len(mismatches)})")
        _table(
            doc,
            ["Severity", "JD requirement", "Gap"],
            [
                [
                    str(m.get("severity") or ""),
                    str(m.get("jd_requirement") or ""),
                    str(m.get("gap") or ""),
                ]
                for m in mismatches
                if isinstance(m, dict)
            ],
        )

    bullets = kit.get("bullets") or []
    if bullets:
        _heading(doc, f"Tailored resume bullets ({len(bullets)})")
        for b in bullets:
            # MatchResult.tailored_bullets is List[Union[str, Dict]] — the LLM has
            # returned both shapes. job_detail.html renders the str case only.
            text = (b.get("bullet") or b.get("text") or "") if isinstance(b, dict) else str(b)
            if text.strip():
                doc.add_paragraph(text.strip(), style="List Bullet")

    hooks = kit.get("hooks") or []
    if hooks:
        _heading(doc, f"Cover-letter hooks ({len(hooks)})")
        for h in hooks:
            text = (h.get("hook") or h.get("text") or "") if isinstance(h, dict) else str(h)
            if text.strip():
                doc.add_paragraph(text.strip(), style="List Bullet")

    cl = kit.get("cover_letter") or {}
    body = cl.get("body") or []
    if body:
        doc.add_page_break()
        _heading(doc, "Cover letter", size=14)
        voice = kit.get("voice") or {}
        if voice.get("enabled"):
            _muted(
                doc,
                f"Voice pass on — {voice.get('before', 0)} flagged phrase(s) before, "
                f"{voice.get('after', 0)} after"
                + (" (rewritten)" if voice.get("rewritten") else ""),
            )
        else:
            _muted(doc, "Voice pass disabled — this text is unpolished LLM output.")

        if cl.get("recipient"):
            doc.add_paragraph(cl["recipient"])
        p = doc.add_paragraph(cl.get("salutation") or "Dear Hiring Team,")
        p.paragraph_format.space_before = Pt(10)
        for para in body:
            if isinstance(para, str) and para.strip():
                bp = doc.add_paragraph(para.strip())
                bp.alignment = WD_ALIGN_PARAGRAPH.LEFT
                bp.paragraph_format.space_after = Pt(8)
        doc.add_paragraph(cl.get("closing") or "Sincerely,")
    elif cl.get("error"):
        _heading(doc, "Cover letter")
        _muted(doc, f"Generation failed: {cl['error']}", size=9)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
