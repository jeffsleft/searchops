"""
Generate the example Accomplishments Inventory used by the demo/fork path (W2).

Why this exists
---------------
`data/Accomplishments_Inventory.docx` is the real Layer 2 input and is gitignored
(`.gitignore`: `data/*.docx`). A fresh clone therefore has no corpus, `corpus.load_corpus()
returns available=False, and Layer 2 contributes 0.0 — which empties the three headline
sections of the Application Kit (evidence, tailored bullets, cover-letter hooks).

This script emits a committed, fictional inventory for the example persona so a stranger
who clones the repo gets a kit with *real* Layer 2 output, produced by the real pipeline,
rather than canned fixtures.

Persona is not invented here — it is fixed by `candidate_profile.example.yaml`:
Alex Rivera, Head of Revenue Operations at AcmeSaaS Inc., $30M→$150M ARR, CPA + RevOps,
seat→consumption pricing transitions, NRR 95%→112%, AI-native builder. Keep this file and
that YAML in sync; if they drift, the demo contradicts its own profile.

Output
------
    data/Accomplishments_Inventory.example.docx

Regenerate with:
    python3 scripts/build_example_corpus.py

Structure contract (enforced by `app/scoring/corpus.py` — do not "improve" casually):
  * `Heading 1` "Narrative Corpus"        → prose block inlined into the Layer 2 prompt
  * `Heading 1` "<Company>  |  <dates>"   → a role section
  * body paragraph under an H1            → that section's context
  * `Heading 2`                           → a theme inside the role
  * 4-column table, header cell 0 containing "Accomplishment" → the accomplishment rows
        columns: Accomplishment | Result | Metric Type | Tags   (Tags are pipe-separated)
  * `Heading 1` "Resume Bullet Swap Library" → H2 theme + `List Paragraph` bullets

Tag vocabulary is deliberately identical to the real inventory's, since the LLM's
`differentiator_themes` are matched against these tags by `corpus.get_bullets_for_themes`:
    ai-native, digital-cs, finance-led, greenfield, health-scoring,
    nrr, resourcefulness, scale, speaking, systems-thinking
"""
from pathlib import Path

from docx import Document
from docx.shared import Pt

OUT_PATH = (
    Path(__file__).resolve().parent.parent
    / "data"
    / "Accomplishments_Inventory.example.docx"
)

NARRATIVE = """\
Alex Rivera — Head of Revenue Operations, AcmeSaaS Inc. (Austin, TX). Eight years building \
GTM systems at Series B–D SaaS companies, preceded by four years in public accounting as a \
licensed CPA (Texas, 2017). Targeting a Director/VP seat owning GTM or Revenue Operations at \
a 200–500 person, Series B+ company. Available October 2026. Remote preferred; hybrid Austin \
or San Francisco one to two times a month is workable.

The through-line is financial rigor applied to go-to-market systems. Most RevOps leaders come \
up through sales operations and learn finance by proximity; Alex audited revenue recognition \
under ASC 606 before ever touching a CRM, and treats pipeline, forecast, and retention data \
with the same tie-out discipline as a general ledger. In practice this shows up as an \
insistence that GTM numbers reconcile to the financials — at AcmeSaaS the CRM-to-GL revenue \
bridge had never tied until Alex closed it to within 0.3%.

At AcmeSaaS (Series B through D, $30M → $150M ARR, ~$79K average ACV) Alex was the second ops \
hire and now leads a team of nine across analytics, systems, and enablement. The two defining \
pieces of work were a pricing-model transition and a renewal rebuild. The pricing work moved \
the company from per-seat to consumption billing; the decision that mattered was killing the \
original big-bang cutover after modeling showed a $2.3M in-year revenue dip, and rebuilding it \
as a renewal-gated rollout instead — which is slower by design, ~60% of ARR converted over five \
quarters rather than all at once, but repapered no customer mid-term and added no churn. The \
renewal work replaced a 30-day "renewal desk" scramble with a 120-day risk trigger fed by usage \
decay, support sentiment, and sponsor departure, moving net revenue retention from 95% to 112% \
and gross retention from 88% to 94% over six quarters.

Alex is honest about the misses, which is the part most inventories leave out. The first health \
score over-indexed on login frequency and cleanly missed three enterprise churns that were \
quietly failing on breadth-of-use; the rebuild onto depth-and-breadth signals is what took \
false negatives from 44% to 12%, and the post-mortem discipline that came out of it — every \
departure coded against the score that missed it — is now standard.

Alex is an AI-native operator in the build sense rather than the user sense: the internal \
renewal-brief generator (Python + LLM) that cut manual renewal prep from 4.5 hours to 1.8 is \
Alex's own code, running in production and used by every CSM, and it shipped only after a \
120-renewal eval harness measured and remediated its fabrication rate. Forecast accuracy \
(±15% to ±6% over four quarters) came from the forecasting rebuild as a whole — bottoms-up \
commit plus weighted pipeline plus a standing variance review — not from any single tool.

Before AcmeSaaS, at Harbor Analytics (Series A → B), Alex built CS Operations from zero as the \
first ops hire — segmentation, coverage model, health scoring, and the QBR motion — including \
a digital-touch program that lifted long-tail gross retention from 71% to 86% with no added \
headcount.

What a hiring manager should probe: Alex has not carried a marketing-ops or demand-gen number \
directly (partnered with it, never owned it), and the largest team led is nine — a VP seat over \
a 20-plus org would be a genuine step up, not a lateral.

Temperamentally: allergic to unread dashboards and undefined metrics. Runs a quarterly kill \
list that deletes any report or automation with no reader in 90 days. Instituted a metric \
definition registry at AcmeSaaS after finding Sales, CS, and Finance each computing "churn" \
three different ways for the same board meeting. Prefers to change the mechanism rather than \
add a person, and will say so in a room where headcount is the expected answer.
"""

# (H1 section name, context paragraph, [(H2 theme, [(accomplishment, result, metric_type, tags)])])
SECTIONS = [
    (
        "AcmeSaaS, Inc.  |  Feb 2021 – Present",
        "Head of Revenue Operations (second ops hire; team of 9 across analytics, systems, "
        "enablement). Series B → Series D, $30M → $150M ARR, ~1,900 accounts, 340 employees.",
        [
            (
                "Pricing Model Transition (seat → consumption)",
                [
                    (
                        "Ran the per-seat to consumption pricing migration across 1,900 accounts: "
                        "modeled revenue impact by cohort, built usage-metering reconciliation "
                        "between the product event pipeline and Zuora, and gated conversion to each "
                        "account's renewal date to avoid mid-term repapering",
                        "~60% of ARR converted over 5 quarters (renewal-gated by design); net "
                        "revenue impact +$4.1M against a modeled +$3.4M, zero migration-attributed "
                        "churn",
                        "Unit Economics",
                        "finance-led | systems-thinking | scale",
                    ),
                    (
                        "Killed the original big-bang cutover plan after modeling showed a $2.3M "
                        "in-year revenue dip and a Q4 support surge — rebuilt it as a "
                        "renewal-triggered rollout over leadership's initial preference for a "
                        "single date",
                        "Avoided the modeled $2.3M dip; zero incremental support headcount through "
                        "the transition",
                        "Finance-led Planning",
                        "finance-led | resourcefulness",
                    ),
                    (
                        "Built the usage-metering reconciliation between the product event pipeline "
                        "and the billing system, with a daily variance report owned by RevOps",
                        "Caught a 4.2% under-billing gap on consumption accounts in month one; "
                        "$1.6M annualized revenue recovered",
                        "Revenue Growth",
                        "finance-led | systems-thinking",
                    ),
                    (
                        "Rebuilt quota and comp plans for the consumption model after the seat-based "
                        "plans started paying on bookings that never converted to usage",
                        "Commission clawbacks 11% → 2% of payout; plan disputes down to zero in two "
                        "quarters",
                        "Unit Economics",
                        "finance-led | scale",
                    ),
                ],
            ),
            (
                "Renewal Operations Rebuild",
                [
                    (
                        "Replaced a 30-day 'renewal desk' scramble with a 120-day risk trigger fed "
                        "by product usage decay, support sentiment, and executive-sponsor churn",
                        "NRR 95% → 112% and GRR 88% → 94% over six quarters",
                        "NRR",
                        "nrr | health-scoring | systems-thinking",
                    ),
                    (
                        "Re-segmented the book by expansion potential rather than ARR band, then "
                        "moved 40% of CSM capacity out of reactive save motions and into expansion "
                        "plays",
                        "Expansion ARR per CSM +38%; save-motion volume −31%",
                        "Retention",
                        "nrr | scale",
                    ),
                    (
                        "Rebuilt the health score after the first version over-indexed on login "
                        "frequency and missed three enterprise churns failing on breadth-of-use — "
                        "moved it onto depth-and-breadth signals and instituted a churn post-mortem "
                        "that codes every departure against the score that missed it",
                        "Health-score false-negative rate 44% → 12% over four quarters; the "
                        "post-mortem is now standard practice",
                        "Churn Risk",
                        "health-scoring | nrr",
                    ),
                    (
                        "Consolidated three overlapping renewal workflows (Sales-led, CS-led, "
                        "self-serve) into one motion with a single owner per account stage",
                        "Renewal cycle time 38 → 19 days; no account entered a renewal window "
                        "unowned",
                        "Cycle Time",
                        "systems-thinking | nrr",
                    ),
                ],
            ),
            (
                "AI-Native Internal Tooling (built, not bought)",
                [
                    (
                        "Wrote and shipped an internal renewal-brief generator (Python + LLM) that "
                        "assembles usage trend, support history, and contract terms into a one-page "
                        "brief the CSM reviews instead of assembling by hand",
                        "Manual renewal prep 4.5 hrs → 1.8 hrs per renewal (−60%); 100% CSM "
                        "adoption within 5 weeks",
                        "Efficiency",
                        "ai-native | greenfield",
                    ),
                    (
                        "Built an LLM-assisted call-note to CRM field extractor after measuring that "
                        "34% of open opportunities carried stale next-step data",
                        "Stale next-step records 34% → 7%; cleaner inputs fed the forecasting "
                        "rebuild rather than standing in for it",
                        "Data Quality",
                        "ai-native | systems-thinking",
                    ),
                    (
                        "Ran an evaluation harness against the brief generator before rollout — 120 "
                        "historical renewals scored by two CSMs — rather than shipping on vibes",
                        "Fabrication rate measured at 3%, remediated to under 1% before launch",
                        "Output Quality",
                        "ai-native | resourcefulness",
                    ),
                    (
                        "Declined to buy a $180K/yr forecasting tool after a two-week bake-off "
                        "showed it underperformed the existing dbt model on the company's own "
                        "historical data",
                        "$180K/yr avoided; bake-off methodology reused for three later tool "
                        "decisions",
                        "Resourcefulness",
                        "resourcefulness | finance-led",
                    ),
                ],
            ),
            (
                "Forecasting & Board-Level Reporting",
                [
                    (
                        "Rebuilt the forecast on a bottoms-up commit and weighted-pipeline blend "
                        "with a documented weekly variance review, replacing a single "
                        "rep-judgment roll-up",
                        "Forecast accuracy ±15% → ±6% over four quarters; held within ±7% through "
                        "the Series D diligence period",
                        "Forecast Accuracy",
                        "finance-led | systems-thinking",
                    ),
                    (
                        "Owned the GTM section of the Series D data room: cohort unit economics, "
                        "CAC payback, and net dollar retention bridges",
                        "$75M Series D closed; no GTM data findings raised in diligence",
                        "Financial Milestone",
                        "finance-led | scale",
                    ),
                    (
                        "Reconciled the CRM-to-GL revenue bridge that had never tied — 1,100 "
                        "opportunities carried no invoice linkage",
                        "Bridge tied to within 0.3%; six days removed from the monthly close",
                        "Data Quality",
                        "finance-led | systems-thinking",
                    ),
                ],
            ),
            (
                "GTM Data Infrastructure",
                [
                    (
                        "Migrated GTM reporting off 60+ hand-maintained CRM reports onto dbt models "
                        "in Snowflake with Looker as the single reporting surface",
                        "One definition of pipeline across functions; report-request queue 3 weeks "
                        "→ same-day",
                        "Data Infrastructure",
                        "systems-thinking | scale",
                    ),
                    (
                        "Instituted a metric definition registry after finding Sales, CS, and "
                        "Finance each computed 'churn' differently — three numbers, three decks, "
                        "one board meeting",
                        "Single definition adopted across all three functions; board pack "
                        "reconciles to the GL",
                        "Data Quality",
                        "systems-thinking | finance-led",
                    ),
                    (
                        "Ran a quarterly kill list: every recurring report, dashboard, and "
                        "automation with no reader in 90 days gets deleted",
                        "Retired 40% of dashboards; roughly 1 FTE of maintenance freed",
                        "Efficiency",
                        "resourcefulness | systems-thinking",
                    ),
                ],
            ),
            (
                "Team Build & Leadership",
                [
                    (
                        "Grew RevOps from 2 to 9 across analytics, systems, and enablement; hired "
                        "the first dedicated CRM admin after measuring 40% of analyst time going to "
                        "config requests",
                        "Analyst time on config 40% → 9%; zero regretted attrition in three years",
                        "Team Build",
                        "scale | greenfield",
                    ),
                    (
                        "Published the team's intake and prioritization model publicly inside the "
                        "company, including what RevOps would not do",
                        "Ad-hoc request volume −45%; escalations to the CRO on prioritization "
                        "stopped entirely",
                        "Leadership",
                        "systems-thinking | scale",
                    ),
                    (
                        "Presented the consumption-pricing migration methodology at two industry "
                        "RevOps meetups and an internal all-hands",
                        "Methodology adopted by two peer companies in the investor network",
                        "Speaking",
                        "speaking | finance-led",
                    ),
                ],
            ),
        ],
    ),
    (
        "Harbor Analytics  |  Jul 2018 – Jan 2021",
        "Manager → Senior Manager, Customer Success Operations (first ops hire). "
        "Series A → Series B, $6M → $28M ARR, ~900 accounts.",
        [
            (
                "CS Operations (built from zero)",
                [
                    (
                        "Built CS Operations from nothing as the first ops hire: account "
                        "segmentation, coverage model, health scoring, and the QBR motion",
                        "Function built from zero; three-person team by the Series B",
                        "Greenfield",
                        "greenfield | digital-cs",
                    ),
                    (
                        "Replaced a gut-feel health score with a weighted model on adoption depth, "
                        "support burden, and sponsor engagement, validated against 18 months of "
                        "historical churn",
                        "90-day churn prediction precision 31% → 68%",
                        "Health Scoring",
                        "health-scoring | nrr",
                    ),
                    (
                        "Rebuilt the coverage model around renewal risk rather than account size "
                        "after finding the largest accounts were also the healthiest",
                        "Freed 2 CSMs of capacity with no change to headcount or churn",
                        "Scope",
                        "resourcefulness | digital-cs",
                    ),
                ],
            ),
            (
                "Digital CS & Long-Tail Coverage",
                [
                    (
                        "Launched a digital-touch motion for the 600-account long tail that had no "
                        "coverage at all — lifecycle email, in-app guides, and a pooled inbox",
                        "Long-tail GRR 71% → 86% with no added headcount",
                        "Adoption",
                        "digital-cs | scale | resourcefulness",
                    ),
                    (
                        "Systematized onboarding into a 30/60/90 with a named success criterion and "
                        "an owner per stage",
                        "Time-to-value 94 → 41 days; onboarding-attributed churn cut by half",
                        "TTV",
                        "digital-cs | systems-thinking",
                    ),
                ],
            ),
        ],
    ),
    (
        "Whitfield & Cole LLP  |  Sep 2014 – Jun 2018",
        "Audit Associate → Audit Senior. Public accounting; subscription and SaaS client "
        "portfolio. Licensed CPA (Texas, 2017).",
        [
            (
                "(general)",
                [
                    (
                        "Led the ASC 606 revenue-recognition transition for a $40M-ARR subscription "
                        "client, rebuilding contract-level revenue schedules from source contracts",
                        "Clean opinion; no restatement",
                        "Audit",
                        "finance-led",
                    ),
                    (
                        "Audited revenue recognition, deferred revenue, and commission "
                        "capitalization across a portfolio of subscription businesses",
                        "Licensed CPA (Texas, 2017)",
                        "Credential",
                        "finance-led",
                    ),
                    (
                        "Built the firm's reusable 606 workpaper template after doing the third "
                        "transition by hand",
                        "Adopted across the practice; roughly 30 hours saved per engagement",
                        "Methodology",
                        "finance-led | resourcefulness",
                    ),
                ],
            ),
        ],
    ),
]

SWAP_LIBRARY = [
    (
        "For roles emphasizing NRR / Retention",
        [
            "Rebuilt the renewal motion around a 120-day risk trigger fed by usage decay, support "
            "sentiment, and sponsor churn — NRR 95% → 112%, GRR 88% → 94% over six quarters.",
            "Re-segmented a 1,900-account book by expansion potential rather than ARR band, moving "
            "40% of CSM capacity from save motions to expansion plays (+38% expansion ARR per CSM).",
            "Cut 90-day churn-prediction false negatives from 44% to 12% by reconciling every "
            "departure against the health score that missed it.",
        ],
    ),
    (
        "For roles emphasizing AI / Automation",
        [
            "Shipped an internal LLM renewal-brief generator now used by every CSM — manual renewal "
            "prep down 60% (4.5 → 1.8 hrs per renewal), full adoption in five weeks.",
            "Built an LLM call-note extractor that cut stale CRM next-step records from 34% to 7% "
            "and lifted forecast accuracy nine points.",
            "Ran a 120-renewal evaluation harness against an internal LLM tool before rollout, "
            "measuring and remediating fabrication from 3% to under 1%.",
        ],
    ),
    (
        "For roles emphasizing Greenfield / Build",
        [
            "Built CS Operations from zero as the first ops hire at a Series A company — "
            "segmentation, coverage, health scoring, and QBR motion; three-person team by Series B.",
            "Launched a digital-touch program for a 600-account long tail with no prior coverage, "
            "lifting GRR from 71% to 86% with no added headcount.",
            "Stood up RevOps at AcmeSaaS as the second ops hire and grew it to nine across "
            "analytics, systems, and enablement.",
        ],
    ),
    (
        "For roles emphasizing Finance / Board-level",
        [
            "Closed a CRM-to-GL revenue bridge that had never tied — 1,100 unlinked opportunities — "
            "to within 0.3%, removing six days from the monthly close.",
            "Owned the GTM section of a $75M Series D data room (cohort unit economics, CAC "
            "payback, NDR bridges); no GTM findings raised in diligence.",
            "Rebuilt forecasting on a bottoms-up commit and weighted-pipeline blend, taking accuracy "
            "from ±14% to ±4% in two quarters.",
        ],
    ),
    (
        "For roles emphasizing Systems / Scale",
        [
            "Migrated GTM reporting off 60+ hand-maintained CRM reports onto dbt models in Snowflake "
            "with Looker as the single surface — report queue from three weeks to same-day.",
            "Ran the seat-to-consumption pricing migration across 1,900 accounts, sequencing by "
            "renewal date; 87% of ARR migrated in three quarters at +$4.1M net revenue impact.",
            "Instituted a metric definition registry after finding three functions computing churn "
            "three different ways for the same board meeting.",
        ],
    ),
    (
        "For roles emphasizing People Leadership",
        [
            "Grew RevOps from 2 to 9 with zero regretted attrition in three years, hiring against "
            "measured time-sinks rather than headcount intuition.",
            "Published the team's intake and prioritization model company-wide — including what "
            "RevOps would not do — cutting ad-hoc requests 45%.",
            "Killed a leadership-preferred big-bang pricing cutover on the strength of a $2.3M "
            "modeled revenue dip, and carried the room on the alternative.",
        ],
    ),
]


def _add_table(doc, rows: list[tuple[str, str, str, str]]) -> None:
    """Emit the 4-column accomplishment table the corpus parser expects."""
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    header = table.rows[0].cells
    # Cell 0 MUST contain "Accomplishment" — corpus._parse_table gates on it.
    for cell, label in zip(header, ("Accomplishment", "Result", "Metric Type", "Tags")):
        cell.text = ""
        run = cell.paragraphs[0].add_run(label)
        run.bold = True
    for text, result, metric_type, tags in rows:
        cells = table.add_row().cells
        cells[0].text = text
        cells[1].text = result
        cells[2].text = metric_type
        cells[3].text = tags
    doc.add_paragraph()


def build(path: Path) -> None:
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    doc.add_heading("Accomplishments Inventory — Alex Rivera (Example)", level=0)
    doc.add_paragraph(
        "Fictional demo corpus for the SearchOps example profile "
        "(candidate_profile.example.yaml). Every company, metric, and person here is "
        "invented. Generated by scripts/build_example_corpus.py — edit that script, not "
        "this file."
    )

    doc.add_heading("Narrative Corpus", level=1)
    for para in NARRATIVE.strip().split("\n\n"):
        doc.add_paragraph(para.replace("\n", " ").strip())

    for name, context, themes in SECTIONS:
        doc.add_heading(name, level=1)
        doc.add_paragraph(context)
        for theme_name, rows in themes:
            doc.add_heading(theme_name, level=2)
            _add_table(doc, rows)

    doc.add_heading("Resume Bullet Swap Library", level=1)
    for theme_name, bullets in SWAP_LIBRARY:
        doc.add_heading(theme_name, level=2)
        for b in bullets:
            doc.add_paragraph(b, style="List Paragraph")

    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))


if __name__ == "__main__":
    build(OUT_PATH)
    print(f"Wrote {OUT_PATH}")
