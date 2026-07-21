"""
W2 — Application Kit tests.

Three things worth pinning, all of which failed silently in the pre-W2 codebase:

  1. The example corpus actually parses. It's a generated binary committed to the repo;
     if `scripts/build_example_corpus.py` drifts from `app/scoring/corpus.py`'s structure
     contract, Layer 2 goes quietly to 0.0 and the kit renders empty for every forker.
  2. The gate is a gate. Score and the manual ethics flag both have to hold, and a blocked
     kit must not burn an LLM call generating a letter nobody asked for.
  3. The voice pass runs on the kit's cover letter. It used to live in the route, so
     assembling via the service layer would have skipped it while still claiming
     "voice-passed".
"""
import os
import tempfile

# Env must exist before any app import (app.config / app.auth read at import time).
_TMP_DB = tempfile.NamedTemporaryFile(suffix=".db", delete=False).name
os.environ.setdefault("SESSION_SECRET", "0" * 64)
os.environ.setdefault("APP_PASSWORD", "test-password")
os.environ.setdefault("GEMINI_API_KEY", "test-key")
os.environ.setdefault("ANTHROPIC_API_KEY", "test-key")
os.environ["DATABASE_PATH"] = _TMP_DB

import json

import pytest

import app.config as config
import app.models as models

config.DATABASE_PATH = _TMP_DB
models.DATABASE_PATH = _TMP_DB

from app.models import get_db, init_db
from app.scoring import corpus as corpus_mod
from app.scoring.corpus import EXAMPLE_INVENTORY_PATH, load_corpus
from app.services.kit_service import build_kit_data, evaluate_gate


@pytest.fixture(scope="module", autouse=True)
def _db():
    init_db()


# ── 1. Example corpus ────────────────────────────────────────────────────────

def test_example_corpus_exists():
    """Committed so a fresh clone has a working Layer 2. Regenerate with the script."""
    assert EXAMPLE_INVENTORY_PATH.exists(), (
        f"{EXAMPLE_INVENTORY_PATH} missing — run: python3 scripts/build_example_corpus.py"
    )


def test_example_corpus_parses_against_the_real_contract():
    c = load_corpus(EXAMPLE_INVENTORY_PATH)
    assert c["available"] is True
    assert c["is_example"] is True
    # A corpus that parses but yields nothing is the failure mode this guards: the parser
    # only reads 4-col tables whose first header cell says "Accomplishment", and silently
    # drops anything else.
    n_bullets = sum(
        len(theme["bullets"]) for s in c["sections"] for theme in s["themes"]
    )
    assert len(c["sections"]) >= 3, "expected the example persona's role sections"
    assert n_bullets >= 20, f"only {n_bullets} accomplishment rows parsed — table contract broken?"
    assert len(c["narrative"]) > 1000, "narrative block is what actually reaches the prompt"
    assert len(c["swap_library"]) == 6


def test_example_corpus_rows_are_fully_populated():
    """Empty result/tag cells mean the table shifted — the parser wouldn't complain."""
    c = load_corpus(EXAMPLE_INVENTORY_PATH)
    for s in c["sections"]:
        for theme in s["themes"]:
            for b in theme["bullets"]:
                assert b["text"].strip(), f"empty accomplishment in {s['name']}/{theme['name']}"
                assert b["result"].strip(), f"empty result for: {b['text'][:60]}"
                assert b["tags"], f"untagged bullet (unreachable by theme match): {b['text'][:60]}"


def test_example_corpus_uses_the_engines_tag_vocabulary():
    """
    Tags are matched against the LLM's `differentiator_themes` by
    corpus.get_bullets_for_themes. An off-vocabulary tag is dead weight.
    """
    known = {
        "ai-native", "digital-cs", "finance-led", "greenfield", "health-scoring",
        "nrr", "resourcefulness", "scale", "speaking", "systems-thinking",
    }
    c = load_corpus(EXAMPLE_INVENTORY_PATH)
    used = {t for s in c["sections"] for th in s["themes"] for b in th["bullets"] for t in b["tags"]}
    assert used <= known, f"off-vocabulary tags: {used - known}"


def test_example_corpus_renders_into_a_prompt():
    c = load_corpus(EXAMPLE_INVENTORY_PATH)
    text = corpus_mod.render_corpus_for_prompt(c)
    assert "Candidate Background" in text
    assert "AcmeSaaS" in text
    assert len(text) > 2000


def test_real_inventory_wins_over_example(monkeypatch, tmp_path):
    """The example is a fallback, never a substitute."""
    fake_real = tmp_path / "Accomplishments_Inventory.docx"
    fake_real.write_bytes(b"stub")
    monkeypatch.setattr(corpus_mod, "INVENTORY_PATH", fake_real)
    path, is_example = corpus_mod.resolve_inventory_path()
    assert path == fake_real
    assert is_example is False


def test_example_used_when_real_inventory_absent(monkeypatch, tmp_path):
    monkeypatch.setattr(corpus_mod, "INVENTORY_PATH", tmp_path / "nope.docx")
    path, is_example = corpus_mod.resolve_inventory_path()
    assert path == EXAMPLE_INVENTORY_PATH
    assert is_example is True


# ── 2. The gate ──────────────────────────────────────────────────────────────

def _job(**over):
    base = {"id": 1, "final_score": 8.0, "ethics_vetted": 1, "auto_rejected": 0}
    base.update(over)
    return base


def test_gate_clears_on_score_plus_ethics():
    assert evaluate_gate(_job())["cleared"] is True


def test_gate_blocks_without_ethics():
    gate = evaluate_gate(_job(ethics_vetted=0))
    assert gate["cleared"] is False
    ethics = next(c for c in gate["checks"] if c["key"] == "ethics")
    assert ethics["ok"] is False


def test_gate_blocks_below_threshold():
    gate = evaluate_gate(_job(final_score=6.9))
    assert gate["cleared"] is False
    assert next(c for c in gate["checks"] if c["key"] == "score")["ok"] is False


def test_gate_blocks_auto_rejected():
    gate = evaluate_gate(_job(auto_rejected=1, final_score=9.0))
    assert gate["cleared"] is False


def test_gate_names_every_failing_check_not_just_the_first():
    """A blocked kit explains itself; fixing one blocker shouldn't reveal a new one."""
    gate = evaluate_gate(_job(final_score=2.0, ethics_vetted=0))
    failing = [c["key"] for c in gate["checks"] if not c["ok"]]
    assert "score" in failing and "ethics" in failing


# ── 3. Assembly + voice pass ─────────────────────────────────────────────────

def _insert_job(**cols) -> int:
    cols.setdefault("date_found", "2026-07-16")
    fields = ", ".join(cols.keys())
    placeholders = ", ".join("?" for _ in cols)
    with get_db() as conn:
        cur = conn.execute(
            f"INSERT INTO jobs ({fields}) VALUES ({placeholders})", tuple(cols.values())
        )
        return cur.lastrowid


@pytest.fixture
def seeded_job():
    """A gate-cleared job row carrying a full Layer 2 payload."""
    return _insert_job(
        company="Northwind Data",
        job_title="Director RevOps",
        final_score=8.2,
        ethics_vetted=1,
        auto_rejected=0,
        recommended_angle="Lead with pricing.",
        match_evidence_json=json.dumps([{
            "jd_requirement": "Own forecasting",
            "matched_accomplishment": "Rebuilt forecast model",
            "strength": "Strong",
        }]),
        match_mismatches_json=json.dumps([
            {"jd_requirement": "CPQ", "gap": "No CPQ", "severity": "High"}
        ]),
        match_bullets_json=json.dumps(["Rebuilt the renewal motion — NRR 95% to 112%."]),
        match_hooks_json=json.dumps(["Your consumption transition is the whole job."]),
    )


def test_build_kit_assembles_every_component(seeded_job):
    kit = build_kit_data(
        seeded_job,
        enrich_fn=lambda r: dict(r),
        cover_letter_loader=lambda job: {"body": ["A paragraph."], "voice": {"enabled": True}},
    )
    assert kit["gate"]["cleared"] is True
    assert kit["angle"] == "Lead with pricing."
    assert kit["evidence"][0]["matched_accomplishment"] == "Rebuilt forecast model"
    assert kit["mismatches"][0]["severity"] == "High"
    assert kit["bullets"] and kit["hooks"]
    assert kit["ready"] is True


def test_blocked_kit_never_generates_a_cover_letter():
    """The LLM call is the expensive part; a gated job must not trigger it."""
    job_id = _insert_job(
        company="Acme", job_title="Ops", final_score=8.5, ethics_vetted=0  # ethics NOT vetted
    )

    calls = []

    def loader(job):
        calls.append(job)
        return {"body": ["should never happen"]}

    kit = build_kit_data(job_id, enrich_fn=lambda r: dict(r), cover_letter_loader=loader)
    assert kit["gate"]["cleared"] is False
    assert kit["cover_letter"] is None
    assert calls == [], "cover letter generated for a job that never cleared the gate"


def test_cover_letter_failure_does_not_sink_the_kit(seeded_job):
    def boom(job):
        raise RuntimeError("provider down")

    kit = build_kit_data(seeded_job, enrich_fn=lambda r: dict(r), cover_letter_loader=boom)
    assert kit["ready"] is False
    assert "provider down" in kit["cover_letter"]["error"]
    assert kit["evidence"], "evidence is still worth showing when the letter fails"


def test_missing_job_returns_none(seeded_job):
    assert build_kit_data(999, lambda r: dict(r), lambda j: {}) is None


# ── 4. Route + template (the only thing that proves kit.html renders) ────────

@pytest.fixture
def client():
    from starlette.testclient import TestClient

    from app.auth import SESSION_COOKIE, create_session_token
    from app.routes import create_app

    c = TestClient(create_app())
    c.cookies.set(SESSION_COOKIE, create_session_token())
    return c


def test_kit_route_renders_for_a_cleared_job(client, seeded_job, monkeypatch):
    import app.routes as routes

    monkeypatch.setattr(
        routes, "_load_cover_letter",
        lambda job: {"salutation": "Dear Hiring Team,", "body": ["A paragraph."],
                     "closing": "Sincerely,", "voice": {"enabled": True, "after": 0}},
    )
    r = client.get(f"/job/{seeded_job}/kit")
    assert r.status_code == 200
    body = r.text
    assert "Northwind Data" in body
    assert "Lead with pricing." in body
    assert "Rebuilt forecast model" in body       # evidence
    assert "No CPQ" in body                        # mismatch
    assert "A paragraph." in body                  # cover letter
    assert "Kit brief .docx" in body
    assert "{{" not in body, "unrendered Jinja left in the page"


def test_kit_route_blocks_and_explains(client, monkeypatch):
    import app.routes as routes

    calls = []
    monkeypatch.setattr(routes, "_load_cover_letter",
                        lambda job: calls.append(job) or {"body": []})

    job_id = _insert_job(company="Blocked Co", job_title="Ops", final_score=8.5,
                         ethics_vetted=0)
    r = client.get(f"/job/{job_id}/kit")
    assert r.status_code == 200
    assert "Mark ethics vetted" in r.text
    assert "gated until every check passes" in r.text
    assert calls == [], "blocked kit route generated a cover letter"


def test_kit_route_404s_for_missing_job(client):
    assert client.get("/job/99999/kit").status_code == 404


def test_kit_download_renders_a_real_docx(client, seeded_job, monkeypatch):
    import app.routes as routes

    monkeypatch.setattr(
        routes, "_load_cover_letter",
        lambda job: {"salutation": "Dear Hiring Team,", "body": ["A paragraph."],
                     "closing": "Sincerely,", "voice": {"enabled": True, "before": 1, "after": 0}},
    )
    r = client.get(f"/job/{seeded_job}/kit/download")
    assert r.status_code == 200
    assert "wordprocessingml" in r.headers["content-type"]
    assert "Northwind Data" in r.headers["content-disposition"]

    # Round-trip it: bytes that aren't openable aren't a document.
    import io

    from docx import Document

    doc = Document(io.BytesIO(r.content))
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "Director RevOps" in text
    assert "Lead with pricing." in text
    table_text = " ".join(
        c.text for t in doc.tables for row in t.rows for c in row.cells
    )
    assert "Rebuilt forecast model" in table_text
    assert "No CPQ" in table_text


def test_kit_download_refuses_a_blocked_job(client):
    job_id = _insert_job(company="Blocked Co 2", job_title="Ops", final_score=2.0,
                         ethics_vetted=0)
    r = client.get(f"/job/{job_id}/kit/download", follow_redirects=False)
    assert r.status_code == 303
    assert r.headers["location"].endswith(f"/job/{job_id}/kit")


def test_generate_cover_letter_applies_the_voice_pass(monkeypatch):
    """
    Regression guard for the W2 move: the polish used to live in routes.py, so anything
    assembling via the service layer got an unpolished letter while the UI claimed
    "voice-passed".
    """
    from app.scoring import cover_letter as cl_mod

    monkeypatch.setattr(cl_mod, "load_corpus", lambda: {"available": True, "narrative": "x",
                                                        "sections": [], "swap_library": {}})
    monkeypatch.setattr(cl_mod, "render_corpus_for_prompt", lambda c: "corpus")
    monkeypatch.setattr(cl_mod, "load_profile", lambda: {})
    monkeypatch.setattr(cl_mod, "_candidate_summary", lambda p: "summary")
    monkeypatch.setattr(cl_mod, "_research_summary", lambda c: "")
    monkeypatch.setattr(cl_mod, "_match_signal", lambda j: "")

    class P:
        def generate_json(self, prompt):
            return {"recipient": "Hiring Team", "salutation": "Dear Hiring Team,",
                    "body": ["We are thrilled to leverage synergies."], "closing": "Sincerely,"}

    monkeypatch.setattr(cl_mod, "get_provider", lambda: P())

    seen = {}

    def fake_polish(cl):
        seen["called"] = True
        cl["voice"] = {"enabled": True, "before": 2, "after": 0, "rewritten": True}
        return cl

    monkeypatch.setattr("app.voice.polish_cover_letter", fake_polish)

    result = cl_mod.generate_cover_letter({"company": "Acme", "jd_text": "jd"})
    assert seen.get("called") is True, "voice pass skipped in generate_cover_letter"
    assert result["voice"]["enabled"] is True
