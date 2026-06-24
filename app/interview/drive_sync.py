import io
from datetime import datetime
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseDownload
from docx import Document
from app.config import INTERVIEW_PREP_DOC_ID

def append_questions_to_prep_doc(company_name, job_title, questions):
    """
    Downloads the Google Doc, appends a new section with questions, and uploads it.
    Uses the Drive API directly (requires google-token-file secret).
    """
    if not INTERVIEW_PREP_DOC_ID:
        print("No INTERVIEW_PREP_DOC_ID configured.")
        return False

    # Initialize Drive service
    # Note: In Modal, we assume the credentials/token are handled by the environment/secret
    # which is already setup for the Sheets API.
    from app.sheets.sync import get_google_creds
    creds = get_google_creds()
    service = build('drive', 'v3', credentials=creds)

    doc_id = INTERVIEW_PREP_DOC_ID
    
    # 1. Download the document as a .docx file
    request = service.files().export_media(fileId=doc_id, mimeType='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    fh = io.BytesIO()
    downloader = MediaIoBaseDownload(fh, request)
    done = False
    while done is False:
        status, done = downloader.next_chunk()

    # 2. Modify the document with python-docx
    fh.seek(0)
    doc = Document(fh)
    
    # Add a heading
    doc.add_page_break()
    doc.add_heading(f"{company_name} — {job_title}", level=1)
    doc.add_paragraph(f"Questions Generated: {datetime.now().strftime('%Y-%m-%d')}")
    
    doc.add_heading("AI-Generated Mock Questions", level=2)
    for q in questions:
        # Assuming questions is a list of dicts from MOCK_QUESTIONS_PROMPT
        # or a list of strings. Let's handle both.
        if isinstance(q, dict):
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(f"{q.get('question', '')}").bold = True
            if q.get('anchor_story'):
                doc.add_paragraph(f"   Anchor: {q['anchor_story']}", style='Normal')
            if q.get('key_point'):
                doc.add_paragraph(f"   Key Point: {q['key_point']}", style='Normal')
        else:
            doc.add_paragraph(str(q), style='List Bullet')

    # 3. Save modified doc to a byte stream
    output_fh = io.BytesIO()
    doc.save(output_fh)
    output_fh.seek(0)

    # 4. Update the existing file on Google Drive
    from googleapiclient.http import MediaIoBaseUpload
    media_body = MediaIoBaseUpload(output_fh, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document', resumable=True)
    
    service.files().update(
        fileId=doc_id,
        media_body=media_body
    ).execute()

    print(f"Successfully appended questions for {company_name} to Google Doc.")
    return True
