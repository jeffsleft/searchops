"""
Generate the locked .docx masters used by the document engine (WP-E).

WHY THIS EXISTS
---------------
The resume/cover-letter generator no longer *builds a Word doc from scratch* in
Python (that was the bug: every render recomputed the design from heuristics, so
formatting drifted whenever content changed). Instead it renders a **locked master
template** with docxtpl, injecting only content.

This script is the "design skill": it is the one place the master templates are
authored. Run it once to (re)generate the two templates the engine renders:

    app/templates/docx/resume_template.docx
    app/templates/docx/cover_letter_template.docx

A user forking the project either:
  * edits this script and re-runs it, or
  * opens the generated .docx in Word and edits fonts/spacing directly (the Jinja
    placeholders like {{ header.name }} and {%p for ... %} loops are visible there).

The templates ship in the repo so the engine works out of the box with no Google
setup and no manual template authoring.

Run:  python scripts/build_doc_templates.py
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

# ── Design tokens (the single source of truth for document design) ──────────
FONT = "Source Serif 4"
BODY_PT = 10.25
LINE_PT = 14  # exact line height; stops Word inflating with its 1.08 multiplier

# US Letter with the margins the on-screen resume_print.html uses
PAGE_W, PAGE_H = 8.5, 11.0
M_LEFT = M_RIGHT = 0.65
M_TOP = 0.55
M_BOT = 0.5
# Right-aligned tab stop position (twips) = content width
CONTENT_W = int((PAGE_W - M_LEFT - M_RIGHT) * 1440)

OUT_DIR = Path(__file__).resolve().parent.parent / "app" / "templates" / "docx"


# ── Low-level docx helpers (mirrors the proven helpers from the old builder) ─
def _font(run, size_pt: float, *, bold=False, italic=False, hex_color: str | None = None):
    run.font.name = FONT
    run.font.size = Pt(size_pt)
    if bold:
        run.font.bold = True
    if italic:
        run.font.italic = True
    if hex_color:
        run.font.color.rgb = RGBColor(
            int(hex_color[0:2], 16), int(hex_color[2:4], 16), int(hex_color[4:6], 16)
        )


def _char_spacing(run, spacing_pt: float):
    rPr = run._r.get_or_add_rPr()
    el = OxmlElement("w:spacing")
    el.set(qn("w:val"), str(int(spacing_pt * 20)))
    rPr.append(el)


def _bottom_border(para, color="555555", sz=4):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(sz))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def _top_border(para, color, sz, space_before, space_after):
    para.paragraph_format.space_before = Pt(space_before)
    para.paragraph_format.space_after = Pt(space_after)
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    top = OxmlElement("w:top")
    top.set(qn("w:val"), "single")
    top.set(qn("w:sz"), str(sz))
    top.set(qn("w:space"), "1")
    top.set(qn("w:color"), color)
    pBdr.append(top)
    pPr.append(pBdr)


def _right_tab(para):
    pPr = para._p.get_or_add_pPr()
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "right")
    tab.set(qn("w:pos"), str(CONTENT_W))
    tabs.append(tab)
    pPr.append(tabs)


def _tag(doc, text: str):
    """A bare Jinja control-flow paragraph (e.g. '{%p for x in xs %}')."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = Pt(0.1)  # collapse the now-empty tag line
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    p.add_run(text)
    return p


def _base_doc() -> Document:
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Inches(PAGE_W)
    sec.page_height = Inches(PAGE_H)
    sec.left_margin = Inches(M_LEFT)
    sec.right_margin = Inches(M_RIGHT)
    sec.top_margin = Inches(M_TOP)
    sec.bottom_margin = Inches(M_BOT)

    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(BODY_PT)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.line_spacing = Pt(LINE_PT)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    return doc


# ── Resume master ────────────────────────────────────────────────────────────
def build_resume_template(path: Path):
    """
    Render context (supplied by app/resume_docx.py):
      header: {name, tagline, contact}
      blocks: [
        {
          first: bool,                # suppress the inter-section rule on the first
          heading: str | "",          # prose section heading (already upper-cased)
          org_row: RichText | None,   # experience: "Company \t Dates"
          role_row: RichText | None,  # experience: "Role \t Location"
          lines:   [RichText, ...],   # prose lines OR experience meta/summary lines
          bullets: [RichText, ...],   # bullets (lead-in bolding pre-applied)
        }, ...
      ]
    Inline-formatted content arrives as RichText so the engine controls bold
    lead-ins; the template owns layout (page, margins, spacing, tab stops, rules).
    """
    doc = _base_doc()

    # Header — name (centered, letter-spaced, bold)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("{{ header.name }}")
    _font(r, 22, bold=True, hex_color="0a0a0a")
    _char_spacing(r, 6.16)

    _tag(doc, "{%p if header.tagline %}")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    _font(p.add_run("{{ header.tagline }}"), 10.5, italic=True, hex_color="333333")
    _tag(doc, "{%p endif %}")

    _tag(doc, "{%p if header.contact %}")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    _font(p.add_run("{{ header.contact }}"), 9.5, hex_color="333333")
    _tag(doc, "{%p endif %}")

    # Header rule (full-width, dark)
    _top_border(doc.add_paragraph(), "111111", 8, 4, 3)

    # Sections
    _tag(doc, "{%p for block in blocks %}")

    _tag(doc, "{%p if not block.first %}")
    _top_border(doc.add_paragraph(), "cccccc", 4, 3, 0)
    _tag(doc, "{%p endif %}")

    # Prose heading (bordered caps)
    _tag(doc, "{%p if block.heading %}")
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)
    _bottom_border(p)
    r = p.add_run("{{ block.heading }}")
    _font(r, 10, bold=True, hex_color="0a0a0a")
    _char_spacing(r, 2.2)
    _tag(doc, "{%p endif %}")

    # Experience org/date row (right-tabbed)
    _tag(doc, "{%p if block.org_row %}")
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(0)
    _right_tab(p)
    p.add_run("{{r block.org_row }}")
    _tag(doc, "{%p endif %}")

    # Experience role/location row (right-tabbed)
    _tag(doc, "{%p if block.role_row %}")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    _right_tab(p)
    p.add_run("{{r block.role_row }}")
    _tag(doc, "{%p endif %}")

    # Body lines (prose lines or experience meta/summary) — RichText carries font
    _tag(doc, "{%p for line in block.lines %}")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.add_run("{{r line }}")
    _tag(doc, "{%p endfor %}")

    # Bullets (hanging indent; lead-in bold pre-applied in RichText)
    _tag(doc, "{%p for bullet in block.bullets %}")
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.first_line_indent = Inches(-0.13)
    p.paragraph_format.space_after = Pt(1)
    p.add_run("{{r bullet }}")
    _tag(doc, "{%p endfor %}")

    _tag(doc, "{%p endfor %}")

    doc.save(str(path))


# ── Cover-letter master ───────────────────────────────────────────────────────
def build_cover_letter_template(path: Path):
    """
    Render context (supplied by app/resume_docx.py):
      header:    {name, contact}
      date:      "June 20, 2026"
      recipient: "Hiring Team, Acme Inc."   (one or more lines via \n)
      salutation:"Dear Hiring Team,"
      body:      [str, ...]              full letter paragraphs (full sentences)
      closing:   "Sincerely,"
      signature: "Jeff Beaumont"
    """
    doc = _base_doc()
    # Letters read better with a little more air than a resume.
    normal = doc.styles["Normal"]
    normal.paragraph_format.line_spacing = Pt(15)

    # Sender name + contact (top of letter)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    _font(p.add_run("{{ header.name }}"), 13, bold=True, hex_color="0a0a0a")

    _tag(doc, "{%p if header.contact %}")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    _font(p.add_run("{{ header.contact }}"), 9.5, hex_color="333333")
    _tag(doc, "{%p endif %}")

    # Date
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    _font(p.add_run("{{ date }}"), BODY_PT)

    # Recipient block
    _tag(doc, "{%p if recipient %}")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    _font(p.add_run("{{ recipient }}"), BODY_PT)
    _tag(doc, "{%p endif %}")

    # Salutation
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    _font(p.add_run("{{ salutation }}"), BODY_PT)

    # Body paragraphs
    _tag(doc, "{%p for para in body %}")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    _font(p.add_run("{{ para }}"), BODY_PT, hex_color="1d1d1d")
    _tag(doc, "{%p endfor %}")

    # Closing + signature
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    _font(p.add_run("{{ closing }}"), BODY_PT)

    p = doc.add_paragraph()
    _font(p.add_run("{{ signature }}"), BODY_PT, bold=True)

    doc.save(str(path))


def main():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    resume = OUT_DIR / "resume_template.docx"
    cover = OUT_DIR / "cover_letter_template.docx"
    build_resume_template(resume)
    build_cover_letter_template(cover)
    print(f"[build_doc_templates] wrote {resume}")
    print(f"[build_doc_templates] wrote {cover}")


if __name__ == "__main__":
    main()
